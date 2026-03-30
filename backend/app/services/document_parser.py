"""
Document parsing utilities for extracting text from various file formats.

Supports:
- PDF files using PyPDF2
- DOCX files using python-docx
- Plain text files (TXT, MD)
"""

import io
import logging
from typing import Dict, List, Optional, Tuple, Union

from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parser for extracting text from different document formats."""

    @staticmethod
    def _sanitize_metadata(metadata: Dict) -> Dict:
        """
        Convert metadata values to JSON-serializable types.
        
        Args:
            metadata: Dictionary potentially containing non-serializable values
            
        Returns:
            Sanitized dictionary with stringified values
        """
        import json
        from datetime import datetime, date
        
        def _serialize(obj):
            if isinstance(obj, (datetime, date)):
                return obj.isoformat()
            elif hasattr(obj, '__dict__'):
                return str(obj)
            else:
                return obj
        
        sanitized = {}
        for key, value in metadata.items():
            try:
                # Try to serialize the value to see if it's JSON-serializable
                json.dumps(value)
                sanitized[key] = value
            except (TypeError, ValueError):
                # If not serializable, convert to string
                sanitized[key] = str(value)
        return sanitized

    @staticmethod
    def parse_pdf(file_content: bytes) -> Tuple[str, Dict[str, any]]:
        """
        Extract text from PDF file.
        
        Args:
            file_content: Raw PDF file bytes
            
        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        try:
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            
            text_parts = []
            # Convert pdf_metadata to serializable dict
            raw_metadata = reader.metadata or {}
            pdf_metadata = DocumentParser._sanitize_metadata(dict(raw_metadata))
            metadata = {
                "page_count": len(reader.pages),
                "pdf_metadata": pdf_metadata,
                "has_encryption": reader.is_encrypted,
            }
            
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {page_num} ---\n{page_text}")
                except Exception as page_error:
                    logger.warning(f"Error extracting text from page {page_num}: {page_error}")
                    text_parts.append(f"--- Page {page_num} ---\n[Error extracting text]")
            
            full_text = "\n\n".join(text_parts)
            
            # Extract metadata if available
            if reader.metadata:
                metadata.update({
                    "author": getattr(reader.metadata, "author", None),
                    "creator": getattr(reader.metadata, "creator", None),
                    "producer": getattr(reader.metadata, "producer", None),
                    "subject": getattr(reader.metadata, "subject", None),
                    "title": getattr(reader.metadata, "title", None),
                    "creation_date": getattr(reader.metadata, "creation_date", None),
                    "modification_date": getattr(reader.metadata, "modification_date", None),
                })

            # Sanitize all metadata to ensure JSON serializability
            metadata = DocumentParser._sanitize_metadata(metadata)
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Failed to parse PDF: {e}")
            raise ValueError(f"Failed to parse PDF file: {str(e)}")

    @staticmethod
    def parse_docx(file_content: bytes) -> Tuple[str, Dict[str, any]]:
        """
        Extract text from DOCX file.
        
        Args:
            file_content: Raw DOCX file bytes
            
        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table_num, table in enumerate(doc.tables, 1):
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        table_text.append(row_text)
                if table_text:
                    text_parts.append(f"\n--- Table {table_num} ---\n" + "\n".join(table_text))
            
            full_text = "\n\n".join(text_parts)
            
            # Extract core properties
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "author": doc.core_properties.author if hasattr(doc.core_properties, "author") else None,
                "title": doc.core_properties.title if hasattr(doc.core_properties, "title") else None,
                "subject": doc.core_properties.subject if hasattr(doc.core_properties, "subject") else None,
                "created": str(doc.core_properties.created) if hasattr(doc.core_properties, "created") else None,
                "modified": str(doc.core_properties.modified) if hasattr(doc.core_properties, "modified") else None,
            }
            
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {e}")
            raise ValueError(f"Failed to parse DOCX file: {str(e)}")

    @staticmethod
    def parse_text(file_content: bytes, encoding: str = "utf-8") -> Tuple[str, Dict[str, any]]:
        """
        Extract text from plain text file.
        
        Args:
            file_content: Raw text file bytes
            encoding: Text encoding to use
            
        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        try:
            # Try to decode with specified encoding
            text = file_content.decode(encoding)
            
            # Try other common encodings if utf-8 fails
        except UnicodeDecodeError:
            try:
                text = file_content.decode("latin-1")
                encoding = "latin-1"
            except UnicodeDecodeError:
                text = file_content.decode("utf-8", errors="ignore")
                encoding = "utf-8 (with errors ignored)"
        
        metadata = {
            "encoding": encoding,
            "character_count": len(text),
            "line_count": text.count("\n") + 1,
        }
        
        return text, metadata

    @staticmethod
    def parse_markdown(file_content: bytes) -> Tuple[str, Dict[str, any]]:
        """
        Extract text from Markdown file.
        
        Args:
            file_content: Raw markdown file bytes
            
        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        try:
            text = file_content.decode("utf-8")
            
            # Count markdown elements
            lines = text.split("\n")
            heading_count = sum(1 for line in lines if line.strip().startswith("#"))
            code_block_count = text.count("```")
            link_count = text.count("[")  # Rough count of markdown links
            
            metadata = {
                "encoding": "utf-8",
                "character_count": len(text),
                "line_count": len(lines),
                "heading_count": heading_count,
                "code_block_count": code_block_count,
                "link_count": link_count,
            }
            
            return text, metadata
            
        except UnicodeDecodeError:
            # Fall back to regular text parsing
            return DocumentParser.parse_text(file_content)

    @classmethod
    def parse_document(cls, file_content: bytes, file_type: str, filename: str = "") -> Tuple[str, Dict[str, any]]:
        """
        Parse document based on file type.
        
        Args:
            file_content: Raw file bytes
            file_type: File extension (pdf, docx, txt, md)
            filename: Original filename (for logging)
            
        Returns:
            Tuple of (extracted_text, metadata_dict)
            
        Raises:
            ValueError: If file type is not supported or parsing fails
        """
        logger.info(f"Parsing document: {filename} (type: {file_type})")
        
        file_type = file_type.lower()
        
        if file_type == "pdf":
            return cls.parse_pdf(file_content)
        elif file_type == "docx":
            return cls.parse_docx(file_content)
        elif file_type == "md":
            return cls.parse_markdown(file_content)
        elif file_type == "txt":
            return cls.parse_text(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")